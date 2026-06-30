import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re

# =========================
# COLORES
# =========================
COLOR_PRINCIPAL="#f57c00"
COLOR_SECUNDARIO="#e65100"
COLOR_FONDO="#fff3e0"
COLOR_PANEL="#ffffff"

referencias=[]

# =========================
# ACERCA DE
# =========================
def mostrar_acerca_de():
    messagebox.showinfo(
        "Acerca de Generador APA 7 PRO",
        "Desarrollado por:\nDaniel García Carbonell\n\nEmail:\nCarbonellgarciadaniel@gmail.com"
    )

# =========================
# UTILIDADES APA
# =========================
def iniciales(nombre):
    partes = nombre.strip().split()
    return " ".join([p[0].upper() + "." for p in partes if p])

def formatear_autor_apa(autor):
    if "," in autor:
        apellido, nombre = autor.split(",", 1)
        return f"{apellido.strip()}, {iniciales(nombre)}"
    return autor

def formatear_lista_autores(autores):
    if not autores:
        return ""
    autores_apa = [formatear_autor_apa(a) for a in autores]
    if len(autores_apa) == 1:
        return autores_apa[0]
    if len(autores_apa) <= 20:
        return ", ".join(autores_apa[:-1]) + " & " + autores_apa[-1]
    primeros = ", ".join(autores_apa[:19])
    return primeros + ", ... " + autores_apa[-1]

def obtener(c,keys):
    for k in keys:
        if c.get(k):
            return c.get(k)
    return ""

def apa_articulo(r):
    # =========================
    # EXTRAER TODOS LOS AUTORES
    # =========================
    autores = []
    for k in r.keys():
        if k.startswith("A") and (k[1:].isdigit() or k[1:]=="U"):
            val = r[k]
            if isinstance(val,list):
                autores.extend(val)
            else:
                autores.append(val)
    # Formatear autores y agregar autor_sort
    autores_formateados = formatear_lista_autores(autores)
    primer_autor_sort = autores[0] if autores else ""
    
    # DOI normalizado
    doi = r.get("DO") or r.get("UR") or ""
    if doi and not doi.startswith("http"):
        doi = "https://doi.org/" + doi

    return{
        "autores": autores_formateados,
        "autor_sort": primer_autor_sort,
        "año": obtener(r,["PY","DP","Y1"]),
        "titulo": obtener(r,["TI","T1"]),
        "revista": obtener(r,["JO","JF","JT"]),
        "doi": doi
    }

# =========================
# WORD
# =========================
def add_hyperlink(paragraph,url,text):
    part=paragraph.part
    r_id=part.relate_to(url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink=OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'),r_id)
    new_run=OxmlElement('w:r')
    text_elem=OxmlElement('w:t')
    text_elem.text=text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def exportar_word():
    if not referencias:
        messagebox.showwarning("Aviso","No hay referencias.")
        return

    # Ordenar por apellido del primer autor
    def apellido_primer_autor(ref):
        autor = ref.get("autor_sort","")
        if "," in autor:
            return autor.split(",")[0].strip().lower()
        else:
            return autor.split()[0].strip().lower()
    refs_ordenadas = sorted(referencias, key=apellido_primer_autor)

    doc=Document()
    doc.add_heading("Referencias",level=1)
    for ref in refs_ordenadas:
        p=doc.add_paragraph()
        p.paragraph_format.left_indent=Inches(0.5)
        p.paragraph_format.first_line_indent=Inches(-0.5)
        p.add_run(f"{ref['autores']} ({ref['año']}). {ref['titulo']}. ")
        if ref["revista"]:
            run=p.add_run(ref["revista"])
            run.italic=True
        if ref["doi"]:
            p.add_run(" ")
            add_hyperlink(p,ref["doi"],ref["doi"])
    doc.save("bibliografia_APA7.docx")
    messagebox.showinfo("Éxito","Word exportado correctamente.")

# =========================
# LECTORES RIS / NBIB
# =========================
def leer_ris(archivo):
    registros=[]
    actual={}
    with open(archivo,encoding="utf-8") as f:
        for linea in f:
            linea=linea.strip()
            if linea.startswith("TY  -"):
                actual={}
            elif linea.startswith("ER  -"):
                registros.append(actual)
            elif "  -" in linea:
                clave,valor=linea.split("  -",1)
                clave=clave.strip()
                valor=valor.strip()
                if clave in actual:
                    if isinstance(actual[clave],list):
                        actual[clave].append(valor)
                    else:
                        actual[clave] = [actual[clave], valor]
                else:
                    actual[clave]=valor
    return registros

def leer_nbib(archivo):
    registros=[]
    actual={}
    with open(archivo,encoding="utf-8") as f:
        for linea in f:
            linea=linea.strip()
            if linea=="":
                if actual:
                    registros.append(actual)
                    actual={}
            elif "-" in linea:
                clave,valor=linea.split("-",1)
                clave=clave.strip()
                valor=valor.strip()
                if clave in actual:
                    if isinstance(actual[clave],list):
                        actual[clave].append(valor)
                    else:
                        actual[clave] = [actual[clave], valor]
                else:
                    actual[clave]=valor
    if actual:
        registros.append(actual)
    return registros

# =========================
# INTERFAZ FUNCIONES
# =========================
def actualizar_lista():
    lista.delete(0,tk.END)
    for r in referencias:
        lista.insert(tk.END,f"{r['titulo']} ({r['año']})")

def cargar_archivos():
    archivos=filedialog.askopenfilenames(
        filetypes=[("RIS/NBIB","*.ris *.nbib")]
    )
    for archivo in archivos:
        ext=os.path.splitext(archivo)[1].lower()
        registros=[]
        if ext==".ris":
            registros=leer_ris(archivo)
        elif ext==".nbib":
            registros=leer_nbib(archivo)
        for r in registros:
            referencias.append(apa_articulo(r))
    actualizar_lista()

def eliminar_referencias():
    sel=lista.curselection()
    if not sel:
        messagebox.showwarning("Eliminar","Selecciona referencias.")
        return
    if not messagebox.askyesno(
        "Confirmar eliminación",
        f"Eliminar {len(sel)} referencias?"
    ):
        return
    for i in reversed(sel):
        del referencias[i]
    actualizar_lista()

# =========================
# INTERFAZ COMERCIAL
# =========================
root=tk.Tk()
root.title("Generador APA 7 PRO")
root.configure(bg=COLOR_FONDO)

try:
    root.iconphoto(False,tk.PhotoImage(file="Icono.png"))
except:
    pass

topbar=tk.Frame(root,bg=COLOR_PRINCIPAL,height=50)
topbar.pack(fill="x")

tk.Label(topbar,
    text="Generador APA 7 PRO",
    bg=COLOR_PRINCIPAL,
    fg="white",
    font=("Segoe UI",14,"bold")
).pack(side="left",padx=15,pady=10)

tk.Button(topbar,
    text="Acerca de",
    bg="white",
    relief="flat",
    command=mostrar_acerca_de
).pack(side="right",padx=10)

panel=tk.Frame(root,bg=COLOR_PANEL)
panel.pack(padx=15,pady=15,fill="both",expand=True)

boton_frame=tk.Frame(panel,bg=COLOR_PANEL)
boton_frame.pack(pady=10)

def boton(txt,cmd,col):
    tk.Button(
        boton_frame,
        text=txt,
        width=25,
        bg=COLOR_PRINCIPAL,
        fg="white",
        relief="flat",
        activebackground=COLOR_SECUNDARIO,
        command=cmd
    ).grid(row=0,column=col,padx=5,pady=5)

boton("Cargar RIS/NBIB",cargar_archivos,0)
boton("Eliminar referencias",eliminar_referencias,1)
boton("Exportar Word",exportar_word,2)

lista=tk.Listbox(panel,
    width=100,
    height=20,
    relief="flat",
    selectmode=tk.EXTENDED
)
lista.pack(pady=10)

root.mainloop()
